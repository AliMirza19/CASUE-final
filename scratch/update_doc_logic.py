
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def update_dashboards_doc(input_path, output_path):
    doc = docx.Document(input_path)
    
    # Data for images 2 to 37
    dashboard_data = [
        {
            "heading": "Admin Dashboard",
            "description": "The Admin Dashboard provides a centralized overview for the System Administrator to manage the CAUSE system. It includes quick access to core management functions like academic terms, user accounts, and budget allocation, ensuring efficient system-wide administration.",
            "caption": "Figure 2.2: CAUSE Society Admin Dashboard Overview"
        },
        {
            "heading": "Budget Management",
            "description": "The Budget Management interface allows the Admin to allocate and track term budgets. It features an immutable budget policy once locked, ensuring financial consistency and preventing unauthorized modifications during the active term.",
            "caption": "Figure 2.3: CAUSE Society Budget Management Interface"
        },
        {
            "heading": "Manage Academic Terms",
            "description": "This interface enables the Admin to create, manage, and activate academic terms. It provides a comprehensive view of all past and current terms, allowing for seamless transitions between different academic sessions.",
            "caption": "Figure 2.4: CAUSE Society Manage Academic Terms Interface"
        },
        {
            "heading": "Manage Users",
            "description": "The Manage Users screen provides tools for the Admin to oversee all system participants, including students, faculty, and specialized teams. It allows for role assignments, profile edits, and account management to maintain system integrity.",
            "caption": "Figure 2.5: CAUSE Society Manage Users Interface"
        },
        {
            "heading": "HOD Management",
            "description": "This module allows the Admin to assign and manage the Head of Department (HOD) for the current term. It tracks assignment history and ensures that the leadership role is properly filled to oversee departmental approvals.",
            "caption": "Figure 2.6: CAUSE Society HOD Management Interface"
        },
        {
            "heading": "Add New Student Data",
            "description": "The Student Data entry interface supports both bulk CSV uploads and manual entry. It ensures that student records are accurately captured following the system's template, facilitating easy integration of new society members.",
            "caption": "Figure 2.7: CAUSE Society Add New Student Data Interface"
        },
        {
            "heading": "Add New Faculty Member",
            "description": "Similar to student entry, the Faculty Registration interface allows for the addition of faculty members via bulk upload or manual forms. This ensures that the academic mentors and supervisors are correctly registered within the system.",
            "caption": "Figure 2.8: CAUSE Society Add New Faculty Member Interface"
        },
        {
            "heading": "User Profile",
            "description": "The My Profile section allows users to view and manage their personal and academic details. It provides a secure way for students, faculty, and admins to update their contact information and maintain their digital identity.",
            "caption": "Figure 2.9: CAUSE Society User Profile Interface"
        },
        {
            "heading": "Presidential Dashboard",
            "description": "The Presidential Dashboard serves as the command center for the society's leader. It provides real-time updates on event reviews, task progress, and team lead assignments, enabling effective oversight of all society operations.",
            "caption": "Figure 2.10: CAUSE Society Presidential Dashboard Overview"
        },
        {
            "heading": "Post Announcement",
            "description": "The Post Announcement interface allows leaders to broadcast important news and updates to the entire society. It features a simple form to create titles and descriptions that appear on the global news feed.",
            "caption": "Figure 2.11: CAUSE Society Post Announcement Interface"
        },
        {
            "heading": "Appoint Team Leads",
            "description": "This interface enables the President to appoint or continue team leads for various roles such as Graphic Designers, Photographers, and Documentation teams, ensuring that every department has a designated leader.",
            "caption": "Figure 2.12: CAUSE Society Appoint Team Leads Interface"
        },
        {
            "heading": "Team Tasks Status",
            "description": "The Team Tasks Status view provides a comprehensive look at all assigned tasks across different teams. The President can monitor progress, review submissions, and provide feedback to ensure quality and timely completion.",
            "caption": "Figure 2.13: CAUSE Society Team Tasks Status Interface"
        },
        {
            "heading": "Assign Team Tasks",
            "description": "The Assign Team Tasks interface allows the President to delegate specific duties to society teams and volunteers. It supports setting due dates and providing detailed instructions for each assignment.",
            "caption": "Figure 2.14: CAUSE Society Assign Team Tasks Interface"
        },
        {
            "heading": "Review Events",
            "description": "The Review Events interface allows the President to evaluate event proposals submitted by students. It facilitates a streamlined workflow where events can be approved, rejected, or sent back for revision based on departmental standards.",
            "caption": "Figure 2.15: CAUSE Society Review Events Interface"
        },
        {
            "heading": "Direct Chat Interface",
            "description": "The Direct Chat interface facilitates real-time communication between society members and leaders. It supports private messaging to coordinate tasks, discuss event details, and foster collaboration within the community.",
            "caption": "Figure 2.16: CAUSE Society Direct Chat Interface"
        },
        {
            "heading": "HOD Dashboard",
            "description": "The HOD Dashboard provides the Head of Department with an overview of recent announcements and pending approvals. It serves as the primary interface for administrative oversight at the departmental level.",
            "caption": "Figure 2.17: CAUSE Society HOD Dashboard Overview"
        },
        {
            "heading": "Review Events (HOD)",
            "description": "The HOD Review Events interface allows the Head of Department to monitor budget status and provide final approval for society events, ensuring they align with university policies and financial constraints.",
            "caption": "Figure 2.18: CAUSE Society HOD Review Events Interface"
        },
        {
            "heading": "Post Announcement (HOD)",
            "description": "Head of Departments can also use the Post Announcement tool to share official departmental news and academic updates with the society members, maintaining a clear line of communication.",
            "caption": "Figure 2.19: CAUSE Society HOD Post Announcement Interface"
        },
        {
            "heading": "Patron Management",
            "description": "This interface allows the HOD to assign and manage the society's Patron for the current term. It tracks assignment history and ensures that the advisory role is consistently filled by qualified faculty members.",
            "caption": "Figure 2.20: CAUSE Society Patron Management Interface"
        },
        {
            "heading": "Election Candidate Review",
            "description": "The Election Candidate Review interface allows the HOD and selection committee to finalize presidential candidates. It displays shortlisted students and their vision statements for formal review and approval.",
            "caption": "Figure 2.21: CAUSE Society Election Candidate Review Interface"
        },
        {
            "heading": "Manage Budget (HOD)",
            "description": "The HOD Budget Management view provides a detailed breakdown of the term's financial status, including total allocation, spent funds, and remaining balance, helping the HOD make informed approval decisions.",
            "caption": "Figure 2.22: CAUSE Society HOD Manage Budget Interface"
        },
        {
            "heading": "Financial Analytics",
            "description": "The Financial Analytics dashboard provides visual representations of budget distribution and event spending trends. It helps the HOD and Admin analyze financial performance over the academic term.",
            "caption": "Figure 2.23: CAUSE Society Financial Analytics Interface"
        },
        {
            "heading": "Financial Reports",
            "description": "The Financial Reports section allows for the generation and download of comprehensive financial summaries. It tracks budget utilization and spending patterns for auditing and documentation purposes.",
            "caption": "Figure 2.24: CAUSE Society Financial Reports Interface"
        },
        {
            "heading": "Patron Post Announcement",
            "description": "The Patron can use this interface to broadcast mentorship-related announcements and society-wide updates, ensuring that their guidance reaches all members effectively.",
            "caption": "Figure 2.25: CAUSE Society Patron Post Announcement Interface"
        },
        {
            "heading": "Events Review (Patron)",
            "description": "The Patron's Events Review interface allows the faculty advisor to pre-approve event budget proposals before they are forwarded to the HOD for final authorization.",
            "caption": "Figure 2.26: CAUSE Society Patron Events Review Interface"
        },
        {
            "heading": "Election Control Panel",
            "description": "The Election Control Panel allows the Patron to configure candidate registration and voting timelines. It provides the tools to manage the democratic process for selecting new society leadership.",
            "caption": "Figure 2.27: CAUSE Society Patron Election Control Panel Interface"
        },
        {
            "heading": "My Events (Student)",
            "description": "The My Events section for students provides a tracking system for their submitted event proposals. It shows the real-time status of each request as it moves through the approval workflow.",
            "caption": "Figure 2.28: CAUSE Society Student My Events Interface"
        },
        {
            "heading": "Request New Event",
            "description": "The Request New Event form allows students to propose new activities. It captures essential details such as event titles, descriptions, dates, and expected venues for administrative review.",
            "caption": "Figure 2.29: CAUSE Society Request New Event Interface"
        },
        {
            "heading": "Guest Speaker Details",
            "description": "When proposing an event, students can provide detailed information about guest speakers and faculty mentors, including their designations and profile links to support the proposal.",
            "caption": "Figure 2.30: CAUSE Society Guest Speaker Details Interface"
        },
        {
            "heading": "Event Requirements",
            "description": "The Event Requirements interface allows students to list all items needed for their event, such as refreshments or equipment. This data is used to calculate the estimated budget for review.",
            "caption": "Figure 2.31: CAUSE Society Event Requirements Interface"
        },
        {
            "heading": "Student Dashboard Overview",
            "description": "The Student Dashboard provides a personalized overview of the user's involvement in the society, including their event submissions, volunteer status, and access to the AI assistant.",
            "caption": "Figure 2.32: CAUSE Society Student Dashboard Interface"
        },
        {
            "heading": "AI Virtual Assistant",
            "description": "The CAUSE-AI Virtual Assistant provides 24/7 support to society members. It can answer questions about society procedures, event guidelines, and system navigation to improve user experience.",
            "caption": "Figure 2.33: CAUSE Society AI Virtual Assistant Interface"
        },
        {
            "heading": "Team Lead Dashboard",
            "description": "The Team Lead Dashboard (e.g., for Graphic Designers) displays tasks assigned specifically to their role. It allows leads to submit their work and view feedback from the President.",
            "caption": "Figure 2.34: CAUSE Society Team Lead Dashboard Interface"
        },
        {
            "heading": "Event Graphics Queue",
            "description": "The Event Graphics Queue allows the design team to manage visual assets for approved events. It ensures that all promotional materials are created and uploaded in a timely manner.",
            "caption": "Figure 2.35: CAUSE Society Event Graphics Queue Interface"
        },
        {
            "heading": "Team Lead Direct Chat",
            "description": "Team leads have access to direct chat to communicate with the President and other members, facilitating specialized coordination for technical tasks like design and documentation.",
            "caption": "Figure 2.36: CAUSE Society Team Lead Direct Chat Interface"
        }
    ]

    # We need to find paragraphs with images and insert text around them
    # Since IMAGE 1 is already handled, we start from the 2nd image found.
    
    image_count = 0
    current_dashboard_idx = 0
    
    new_doc = docx.Document()
    # Copy initial content until IMAGE 1's caption
    # Or just iterate and insert
    
    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            # Check if this paragraph contains an image
            contains_image = False
            for run in block.runs:
                if 'graphic' in run.element.xml:
                    contains_image = True
                    break
            
            if contains_image:
                image_count += 1
                if image_count > 1:
                    # This is image 2 or later
                    if current_dashboard_idx < len(dashboard_data):
                        data = dashboard_data[current_dashboard_idx]
                        current_dashboard_idx += 1
                        
                        # Add Heading
                        h = new_doc.add_heading(data['heading'], level=2)
                        
                        # Add Description
                        p_desc = new_doc.add_paragraph(data['description'])
                        
                        # Add Image (we need to copy the image paragraph)
                        # To copy a paragraph with an image is hard, but we can just add a placeholder
                        # and then we will use a different approach: modifying the existing doc.
                        pass
                
                # If we are modifying the existing doc, it's easier to insert before/after
                pass
    
    # NEW STRATEGY: Modify the document in place by iterating through paragraphs
    
    doc = docx.Document(input_path)
    paragraphs = list(doc.paragraphs)
    
    image_paragraphs = []
    for i, p in enumerate(paragraphs):
        for run in p.runs:
            if 'graphic' in run.element.xml:
                image_paragraphs.append(i)
                break
    
    print(f"Found {len(image_paragraphs)} image paragraphs")
    
    # We will work backwards to not mess up indices
    for i in range(len(image_paragraphs)-1, 0, -1): # From last image to 2nd image (index 1)
        img_para_idx = image_paragraphs[i]
        data_idx = i - 1 # image_paragraphs[1] is the 2nd image, maps to dashboard_data[0]
        
        if data_idx < len(dashboard_data):
            data = dashboard_data[data_idx]
            
            # Insert Caption AFTER image
            caption_para = paragraphs[img_para_idx].insert_paragraph_before("") # Temporary
            # Actually, insert_paragraph_before inserts BEFORE.
            # To insert AFTER, we use the next paragraph or add one.
            
            # Let's use a different way to insert after
            # Caption
            caption = doc.paragraphs[img_para_idx].add_run().add_break() # No, this adds to same para
            
            # Best way to insert after is to find the paragraph after the image
            # But what if there isn't one?
            
    # Actually, let's just create a NEW document and copy elements.
    # It's more reliable.
    
    print("Done")

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Unknown parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)

# I'll use a simpler script to just insert text into the document using a loop.
# I'll use `p.insert_paragraph_before(text)` for heading and description.
# And `p.add_run().add_break(); p.add_run(caption)` might not work for a new para.

# Let's write the FINAL implementation script.
