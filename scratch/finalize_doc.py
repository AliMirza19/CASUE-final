
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def update_dashboards():
    input_file = "CAUSE DASHBORDS.docx"
    output_file = "CAUSE DASHBORDS COMPLETE.docx"
    
    doc = docx.Document(input_file)
    
    # Content for images 2 to 37
    dashboard_data = [
        {"h": "Admin Dashboard", "d": "The Admin Dashboard provides a centralized overview for the System Administrator to manage the CAUSE system. It includes quick access to core management functions like academic terms, user accounts, and budget allocation, ensuring efficient system-wide administration.", "c": "Figure 2.2: CAUSE Society Admin Dashboard Overview"},
        {"h": "Budget Management", "d": "The Budget Management interface allows the Admin to allocate and track term budgets. It features an immutable budget policy once locked, ensuring financial consistency and preventing unauthorized modifications during the active term.", "c": "Figure 2.3: CAUSE Society Budget Management Interface"},
        {"h": "Manage Academic Terms", "d": "This interface enables the Admin to create, manage, and activate academic terms. It provides a comprehensive view of all past and current terms, allowing for seamless transitions between different academic sessions.", "c": "Figure 2.4: CAUSE Society Manage Academic Terms Interface"},
        {"h": "Manage Users", "d": "The Manage Users screen provides tools for the Admin to oversee all system participants, including students, faculty, and specialized teams. It allows for role assignments, profile edits, and account management to maintain system integrity.", "c": "Figure 2.5: CAUSE Society Manage Users Interface"},
        {"h": "HOD Management", "d": "This module allows the Admin to assign and manage the Head of Department (HOD) for the current term. It tracks assignment history and ensures that the leadership role is properly filled to oversee departmental approvals.", "c": "Figure 2.6: CAUSE Society HOD Management Interface"},
        {"h": "Add New Student Data", "d": "The Student Data entry interface supports both bulk CSV uploads and manual entry. It ensures that student records are accurately captured following the system's template, facilitating easy integration of new society members.", "c": "Figure 2.7: CAUSE Society Add New Student Data Interface"},
        {"h": "Add New Faculty Member", "d": "Similar to student entry, the Faculty Registration interface allows for the addition of faculty members via bulk upload or manual forms. This ensures that the academic mentors and supervisors are correctly registered within the system.", "c": "Figure 2.8: CAUSE Society Add New Faculty Member Interface"},
        {"h": "User Profile", "d": "The My Profile section allows users to view and manage their personal and academic details. It provides a secure way for students, faculty, and admins to update their contact information and maintain their digital identity.", "c": "Figure 2.9: CAUSE Society User Profile Interface"},
        {"h": "Presidential Dashboard", "d": "The Presidential Dashboard serves as the command center for the society's leader. It provides real-time updates on event reviews, task progress, and team lead assignments, enabling effective oversight of all society operations.", "c": "Figure 2.10: CAUSE Society Presidential Dashboard Overview"},
        {"h": "Post Announcement", "d": "The Post Announcement interface allows leaders to broadcast important news and updates to the entire society. It features a simple form to create titles and descriptions that appear on the global news feed.", "c": "Figure 2.11: CAUSE Society Post Announcement Interface"},
        {"h": "Appoint Team Leads", "d": "This interface enables the President to appoint or continue team leads for various roles such as Graphic Designers, Photographers, and Documentation teams, ensuring that every department has a designated leader.", "c": "Figure 2.12: CAUSE Society Appoint Team Leads Interface"},
        {"h": "Team Tasks Status", "d": "The Team Tasks Status view provides a comprehensive look at all assigned tasks across different teams. The President can monitor progress, review submissions, and provide feedback to ensure quality and timely completion.", "c": "Figure 2.13: CAUSE Society Team Tasks Status Interface"},
        {"h": "Assign Team Tasks", "d": "The Assign Team Tasks interface allows the President to delegate specific duties to society teams and volunteers. It supports setting due dates and providing detailed instructions for each assignment.", "c": "Figure 2.14: CAUSE Society Assign Team Tasks Interface"},
        {"h": "Review Events", "d": "The Review Events interface allows the President to evaluate event proposals submitted by students. It facilitates a streamlined workflow where events can be approved, rejected, or sent back for revision based on departmental standards.", "c": "Figure 2.15: CAUSE Society Review Events Interface"},
        {"h": "Direct Chat Interface", "d": "The Direct Chat interface facilitates real-time communication between society members and leaders. It supports private messaging to coordinate tasks, discuss event details, and foster collaboration within the community.", "c": "Figure 2.16: CAUSE Society Direct Chat Interface"},
        {"h": "HOD Dashboard", "d": "The HOD Dashboard provides the Head of Department with an overview of recent announcements and pending approvals. It serves as the primary interface for administrative oversight at the departmental level.", "c": "Figure 2.17: CAUSE Society HOD Dashboard Overview"},
        {"h": "Review Events (HOD)", "d": "The HOD Review Events interface allows the Head of Department to monitor budget status and provide final approval for society events, ensuring they align with university policies and financial constraints.", "c": "Figure 2.18: CAUSE Society HOD Review Events Interface"},
        {"h": "Post Announcement (HOD)", "d": "Head of Departments can also use the Post Announcement tool to share official departmental news and academic updates with the society members, maintaining a clear line of communication.", "c": "Figure 2.19: CAUSE Society HOD Post Announcement Interface"},
        {"h": "Patron Management", "d": "This interface allows the HOD to assign and manage the society's Patron for the current term. It tracks assignment history and ensures that the advisory role is consistently filled by qualified faculty members.", "c": "Figure 2.20: CAUSE Society Patron Management Interface"},
        {"h": "Election Candidate Review", "d": "The Election Candidate Review interface allows the HOD and selection committee to finalize presidential candidates. It displays shortlisted students and their vision statements for formal review and approval.", "c": "Figure 2.21: CAUSE Society Election Candidate Review Interface"},
        {"h": "Manage Budget (HOD)", "d": "The HOD Budget Management view provides a detailed breakdown of the term's financial status, including total allocation, spent funds, and remaining balance, helping the HOD make informed approval decisions.", "c": "Figure 2.22: CAUSE Society HOD Manage Budget Interface"},
        {"h": "Financial Analytics", "d": "The Financial Analytics dashboard provides visual representations of budget distribution and event spending trends. It helps the HOD and Admin analyze financial performance over the academic term.", "c": "Figure 2.23: CAUSE Society Financial Analytics Interface"},
        {"h": "Financial Reports", "d": "The Financial Reports section allows for the generation and download of comprehensive financial summaries. It tracks budget utilization and spending patterns for auditing and documentation purposes.", "c": "Figure 2.24: CAUSE Society Financial Reports Interface"},
        {"h": "Patron Post Announcement", "d": "The Patron can use this interface to broadcast mentorship-related announcements and society-wide updates, ensuring that their guidance reaches all members effectively.", "c": "Figure 2.25: CAUSE Society Patron Post Announcement Interface"},
        {"h": "Events Review (Patron)", "d": "The Patron's Events Review interface allows the faculty advisor to pre-approve event budget proposals before they are forwarded to the HOD for final authorization.", "c": "Figure 2.26: CAUSE Society Patron Events Review Interface"},
        {"h": "Post Announcement (Patron Duplicate)", "d": "This view shows the post-announcement interface for the Patron role, enabling consistent communication across all leadership levels within the CAUSE system.", "c": "Figure 2.27: CAUSE Society Patron Announcement Management"},
        {"h": "Election Control Panel", "d": "The Election Control Panel allows the Patron to configure candidate registration and voting timelines. It provides the tools to manage the democratic process for selecting new society leadership.", "c": "Figure 2.28: CAUSE Society Patron Election Control Panel Interface"},
        {"h": "My Events (Student)", "d": "The My Events section for students provides a tracking system for their submitted event proposals. It shows the real-time status of each request as it moves through the approval workflow.", "c": "Figure 2.29: CAUSE Society Student My Events Interface"},
        {"h": "Request New Event", "d": "The Request New Event form allows students to propose new activities. It captures essential details such as event titles, descriptions, dates, and expected venues for administrative review.", "c": "Figure 2.30: CAUSE Society Request New Event Interface"},
        {"h": "Guest Speaker Details", "d": "When proposing an event, students can provide detailed information about guest speakers and faculty mentors, including their designations and profile links to support the proposal.", "c": "Figure 2.31: CAUSE Society Guest Speaker Details Interface"},
        {"h": "Event Requirements", "d": "The Event Requirements interface allows students to list all items needed for their event, such as refreshments or equipment. This data is used to calculate the estimated budget for review.", "c": "Figure 2.32: CAUSE Society Event Requirements Interface"},
        {"h": "Student Dashboard Overview", "d": "The Student Dashboard provides a personalized overview of the user's involvement in the society, including their event submissions, volunteer status, and access to the AI assistant.", "c": "Figure 2.33: CAUSE Society Student Dashboard Interface"},
        {"h": "AI Virtual Assistant", "d": "The CAUSE-AI Virtual Assistant provides 24/7 support to society members. It can answer questions about society procedures, event guidelines, and system navigation to improve user experience.", "c": "Figure 2.34: CAUSE Society AI Virtual Assistant Interface"},
        {"h": "Team Lead Dashboard", "d": "The Team Lead Dashboard (e.g., for Graphic Designers) displays tasks assigned specifically to their role. It allows leads to submit their work and view feedback from the President.", "c": "Figure 2.35: CAUSE Society Team Lead Dashboard Interface"},
        {"h": "Event Graphics Queue", "d": "The Event Graphics Queue allows the design team to manage visual assets for approved events. It ensures that all promotional materials are created and uploaded in a timely manner.", "c": "Figure 2.36: CAUSE Society Event Graphics Queue Interface"},
        {"h": "Team Lead Direct Chat", "d": "Team leads have access to direct chat to communicate with the President and other members, facilitating specialized coordination for technical tasks like design and documentation.", "c": "Figure 2.37: CAUSE Society Team Lead Direct Chat Interface"}
    ]
    
    # Identify paragraphs with images
    image_para_indices = []
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            if 'graphic' in run.element.xml:
                image_para_indices.append(i)
                break
    
    print(f"Found {len(image_para_indices)} images.")
    
    # We will iterate backwards to insert without breaking indices
    # But wait, inserting paragraphs shifts the indices of subsequent paragraphs.
    # We can use a different strategy: create a list of instructions and execute from bottom up.
    
    for i in range(len(image_para_indices)-1, 0, -1): # From last image to 2nd image (index 1)
        img_idx = image_para_indices[i]
        data_idx = i - 1 # 2nd image is dashboard_data[0]
        
        if data_idx < len(dashboard_data):
            data = dashboard_data[data_idx]
            
            # Insert Caption AFTER image
            # We insert a paragraph after the image paragraph
            # To insert after, we insert before the next paragraph
            if img_idx + 1 < len(doc.paragraphs):
                p_after = doc.paragraphs[img_idx + 1].insert_paragraph_before(data['c'])
            else:
                p_after = doc.add_paragraph(data['c'])
            p_after.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_after.runs[0].bold = True
            
            # Insert Heading and Description BEFORE image
            # Description
            p_desc = doc.paragraphs[img_idx].insert_paragraph_before(data['d'])
            # Heading
            p_head = p_desc.insert_paragraph_before(data['h'])
            p_head.style = doc.styles['Heading 2']
            
    doc.save(output_file)
    print(f"Saved to {output_file}")

if __name__ == "__main__":
    update_dashboards()
